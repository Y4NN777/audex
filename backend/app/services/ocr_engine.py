"""OCR engine bootstrap layer providing EasyOCR + fallbacks."""

from __future__ import annotations

import queue
import threading
from logging import getLogger
from pathlib import Path
from typing import Protocol, Sequence

try:  # pragma: no cover - optional dependency
    import numpy as np
except Exception:  # noqa: BLE001
    np = None  # type: ignore[assignment]

try:  # pragma: no cover - optional dependency
    import cv2 
except Exception:  # noqa: BLE001
    cv2 = None  # type: ignore[assignment]

try:  # pragma: no cover - optional dependency
    import fitz  # type: ignore[attr-defined]
except Exception:  # noqa: BLE001
    fitz = None  # type: ignore[assignment]

try:  # pragma: no cover - optional dependency
    from docx import Document  # type: ignore[import]
except Exception:  # noqa: BLE001
    Document = None  # type: ignore[assignment]

from PIL import Image

try:  # pragma: no cover - optional dependency
    import easyocr  # type: ignore
except Exception:  # noqa: BLE001
    easyocr = None  # type: ignore[assignment]

from app.core.config import settings
from app.pipelines import ocr as legacy_ocr
from app.pipelines.models import OCRResult
from app.schemas.ingestion import FileMetadata

logger = getLogger(__name__)


class OCREngine(Protocol):
    engine_id: str

    def extract(self, file_meta: FileMetadata) -> OCRResult:
        ...


class LegacyOCREngine:
    """Wrapper around the historical OCR module (pytesseract-based)."""

    engine_id = "tesseract"

    def extract(self, file_meta: FileMetadata) -> OCRResult:
        path = Path(file_meta.stored_path)
        text = ""
        confidence = None
        warnings: list[str] = []

        content_type = (file_meta.content_type or "").lower()

        if content_type.startswith("image/"):
            text = legacy_ocr.extract_text(path)
        elif content_type == "text/plain":
            try:
                text = path.read_text(encoding="utf-8")
                confidence = 1.0 if text else None
            except Exception as exc:  # noqa: BLE001
                warnings.append(f"plain-text-read-error:{exc}")
        elif content_type == "application/pdf":
            text = "[pdf-ocr-unavailable]"
            warnings.append("pdf-ocr-unavailable")
        elif content_type in {
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }:
            text = "[docx-ocr-unavailable]"
            warnings.append("docx-ocr-unavailable")
        else:
            text = ""

        return OCRResult(
            source_file=file_meta.filename,
            text=text.strip(),
            confidence=confidence,
            warnings=warnings,
            error=None,
        )


class EasyOCREngine:
    """Adapter around EasyOCR with PDF/DOCX support and graceful degradation."""

    engine_id = "easyocr"

    def __init__(self, languages: Sequence[str]) -> None:
        self._languages = list(languages) or ["en"]
        self._reader: "easyocr.Reader | None" = None  # type: ignore[name-defined]
        self._lock = threading.Lock()
        self._initialisation_failed = False
        self._init_timeout = max(1, settings.EASY_OCR_INIT_TIMEOUT_SECONDS)

    @staticmethod
    def is_available() -> bool:
        return easyocr is not None

    def extract(self, file_meta: FileMetadata) -> OCRResult:
        path = Path(file_meta.stored_path)
        content_type = (file_meta.content_type or "").lower()

        try:
            if content_type.startswith("image/"):
                return self._extract_image(path, file_meta.filename)
            if content_type == "application/pdf":
                return self._extract_pdf(path, file_meta.filename)
            if content_type in {
                "application/msword",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            }:
                return self._extract_docx(path, file_meta.filename)
            if content_type == "text/plain":
                return self._extract_text_file(path, file_meta.filename)
            return OCRResult(
                source_file=file_meta.filename,
                text="",
                confidence=None,
                warnings=[f"unsupported-content-type:{content_type or 'unknown'}"],
                error=None,
            )
        except Exception as exc:  # noqa: BLE001
            logger.error("OCR pipeline failed on %s: %s", file_meta.filename, exc)
            fallback_text = ""
            if content_type.startswith("image/"):
                fallback_text = legacy_ocr.extract_text(path)
            return OCRResult(
                source_file=file_meta.filename,
                text=fallback_text.strip(),
                confidence=None,
                warnings=[f"ocr-fallback:{self.engine_id}"],
                error=str(exc),
            )

    def _get_reader(self) -> "easyocr.Reader":  # type: ignore[name-defined]
        if self._reader is not None:
            return self._reader
        if self._initialisation_failed:
            raise RuntimeError("EasyOCR initialisation previously failed.")

        if easyocr is None:
            raise RuntimeError("EasyOCR is not available in the current environment.")

        with self._lock:
            if self._reader is None:
                logger.info(
                    "Initialising EasyOCR reader for languages %s (gpu=off).",
                    ", ".join(self._languages),
                )
                result: "queue.Queue[tuple[str, object]]" = queue.Queue()

                def initialise_reader() -> None:
                    try:
                        reader = easyocr.Reader(  # type: ignore[attr-defined]
                            self._languages,
                            gpu=False,
                            download_enabled=True,
                        )
                        result.put(("ok", reader))
                    except Exception as exc:  # noqa: BLE001
                        result.put(("error", exc))

                worker = threading.Thread(target=initialise_reader, name="easyocr-init", daemon=True)
                worker.start()
                worker.join(self._init_timeout)

                if worker.is_alive():
                    self._initialisation_failed = True
                    logger.warning(
                        "EasyOCR initialisation timed out after %ss for languages=%s. Falling back to legacy.",
                        self._init_timeout,
                        ", ".join(self._languages),
                    )
                    raise RuntimeError("EasyOCR initialisation timed out")

                status, payload = result.get()
                if status == "ok":
                    self._reader = payload  # type: ignore[assignment]
                else:
                    self._initialisation_failed = True
                    logger.warning(
                        "Unable to initialise EasyOCR (languages=%s): %s",
                        ", ".join(self._languages),
                        payload,
                    )
                    raise payload  # type: ignore[misc]
        return self._reader

    def _extract_image(self, path: Path, filename: str) -> OCRResult:
        if easyocr is None:
            logger.warning("EasyOCR not available, falling back to legacy for %s", filename)
            text = legacy_ocr.extract_text(path)
            return OCRResult(source_file=filename, text=text.strip(), confidence=None, warnings=["easyocr-missing"])

        image_input = self._prepare_image(path)
        text, confidence = self._read_easyocr(image_input)
        return OCRResult(source_file=filename, text=text, confidence=confidence, warnings=[])

    def _extract_pdf(self, path: Path, filename: str) -> OCRResult:
        warnings: list[str] = []
        collected: list[str] = []
        confidences: list[float] = []

        if fitz is None:
            warnings.append("pymupdf-missing")
            return OCRResult(source_file=filename, text="", confidence=None, warnings=warnings, error="pymupdf-missing")

        with fitz.open(path) as document:  # type: ignore[arg-type]
            for index, page in enumerate(document, start=1):
                page_text = page.get_text().strip()
                if page_text:
                    collected.append(page_text)
                    continue

                if easyocr is None:
                    warnings.append(f"page-{index}:easyocr-missing")
                    continue

                try:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)  # type: ignore[attr-defined]
                    image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    if np is not None:
                        image_input = np.array(image)
                    else:
                        image_input = image
                    text, confidence = self._read_easyocr(image_input)
                    if text:
                        collected.append(text)
                    if confidence is not None:
                        confidences.append(confidence)
                except Exception as exc:  # noqa: BLE001
                    warnings.append(f"page-{index}:ocr-error")
                    logger.warning("OCR on PDF page %s failed (%s): %s", index, filename, exc)

        combined = "\n\n".join(text for text in collected if text).strip()
        if not combined:
            warnings.append("pdf-empty")

        confidence = None
        if confidences:
            confidence = sum(confidences) / len(confidences)
        elif combined:
            confidence = 1.0

        return OCRResult(
            source_file=filename,
            text=combined,
            confidence=confidence,
            warnings=warnings,
        )

    def _extract_docx(self, path: Path, filename: str) -> OCRResult:
        warnings: list[str] = []
        if Document is None:
            warnings.append("python-docx-missing")
            return OCRResult(source_file=filename, text="", confidence=None, warnings=warnings, error="python-docx-missing")

        document = Document(str(path))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        combined = "\n".join(parts).strip()
        confidence = 1.0 if combined else None

        return OCRResult(source_file=filename, text=combined, confidence=confidence, warnings=warnings)

    def _extract_text_file(self, path: Path, filename: str) -> OCRResult:
        try:
            content = path.read_text(encoding="utf-8")
            return OCRResult(source_file=filename, text=content.strip(), confidence=1.0 if content else None, warnings=[])
        except UnicodeDecodeError:
            logger.warning("UTF-8 decoding failed for %s, attempting binary fallback.", filename)
            raw = path.read_bytes()
            return OCRResult(
                source_file=filename,
                text=raw.decode("latin-1", errors="ignore").strip(),
                confidence=None,
                warnings=["text-decoding-latin1"],
            )

    def _prepare_image(self, path: Path) -> object:
        if cv2 is None or np is None:
            return str(path)

        image = cv2.imread(str(path), cv2.IMREAD_COLOR)
        if image is None:
            return str(path)

        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        denoised = cv2.bilateralFilter(gray, 9, 75, 75)
        processed = cv2.adaptiveThreshold(
            denoised,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31,
            2,
        )
        return processed

    def _read_easyocr(self, image_input: object) -> tuple[str, float | None]:
        reader = self._get_reader()
        results = reader.readtext(image_input, detail=1, paragraph=True)  # type: ignore[attr-defined]

        texts: list[str] = []
        confidences: list[float] = []

        for entry in results:
            if not isinstance(entry, (list, tuple)) or len(entry) < 3:
                continue
            _, text, score = entry[:3]
            if isinstance(text, str) and text.strip():
                texts.append(text.strip())
            if isinstance(score, (float, int)):
                confidences.append(float(score))

        combined = "\n".join(texts).strip()
        confidence = None
        if confidences:
            confidence = sum(confidences) / len(confidences)

        return combined, confidence


def get_ocr_engine() -> OCREngine:
    engine = settings.OCR_ENGINE.lower().strip()
    if engine == "easyocr":
        if EasyOCREngine.is_available():
            return EasyOCREngine(settings.OCR_LANGUAGES)
        logger.warning("EasyOCR requested but unavailable; falling back to legacy OCR.")
        return LegacyOCREngine()
    if engine == "tesseract":
        return LegacyOCREngine()
    raise ValueError(f"Unsupported OCR engine: {settings.OCR_ENGINE}")
